import io
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_pdf import MarkdownPdf, Section

def generate_docx(report: Dict[str, Any]) -> bytes:
    """
    Generates a DOCX document from a report dictionary.
    This is a Python port of the `generateDocx` function from `lib/documents.ts`.
    """
    try:
        doc = Document()

        # --- Header ---
        header_section = doc.sections[0]
        header = header_section.header
        header_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_p.text = report.get('title', 'Untitled Report')
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Title ---
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(report.get('title', 'Untitled Report'))
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Summary ---
        summary_p = doc.add_paragraph(report.get('summary', ''))
        summary_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # --- Sections ---
        for section in report.get('sections', []):
            doc.add_heading(section.get('title', 'Untitled Section'), level=1)
            doc.add_paragraph(section.get('content', ''))

        # Save to a byte stream
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        print(f"Error generating DOCX: {e}")
        raise

def generate_pdf(report: Dict[str, Any]) -> bytes:
    """
    Generates a PDF document from a report dictionary using markdown.
    This is a Python port of the `generatePdf` function from `lib/documents.ts`.
    """
    try:
        # Construct a markdown string from the report
        markdown_content = f"# {report.get('title', 'Untitled Report')}\n\n"
        markdown_content += f"**Summary:** {report.get('summary', '')}\n\n"

        for section in report.get('sections', []):
            markdown_content += f"## {section.get('title', 'Untitled Section')}\n\n"
            markdown_content += f"{section.get('content', '')}\n\n"

        # Generate PDF from markdown
        pdf = MarkdownPdf(toc_level=2)
        pdf.add_section(Section(markdown_content, toc=False))

        # The library saves to a file, so we'll use a BytesIO buffer
        buffer = io.BytesIO()
        pdf.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        print(f"Error generating PDF: {e}")
        raise